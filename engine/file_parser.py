"""
TestFortge — Multi-format File Parser

Parses requirements from various file types:
  .txt, .md          — plain text
  .docx              — Microsoft Word (python-docx)
  .xlsx, .csv        — spreadsheets (openpyxl / csv)
  .pdf               — PDF documents (pypdf)
  .png, .jpg, .jpeg  — images (metadata)
  .mp4, .webm, .avi, .mov, .mkv, .flv, .wmv, .gif — video (metadata as attachment)
"""

import csv
import io
import os
import re
from dataclasses import dataclass, field

from engine.log import get_logger

_logger = get_logger(__name__)

# Optional imports — graceful fallback
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from openpyxl import load_workbook
    HAS_XLSX = True
except ImportError:
    HAS_XLSX = False

try:
    # Prefer the actively maintained ``pypdf`` library; fall back to the
    # deprecated ``PyPDF2`` only for legacy environments. Both expose the
    # same ``PdfReader`` class, so downstream code stays untouched.
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:  # pragma: no cover
    try:
        from PyPDF2 import PdfReader  # type: ignore[no-redef]
        HAS_PDF = True
    except ImportError:
        HAS_PDF = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


ALLOWED_EXTENSIONS = {
    # Text / Document
    "txt", "md", "doc", "docx",
    # Spreadsheet
    "xlsx", "csv",
    # PDF
    "pdf",
    # Image
    "png", "jpg", "jpeg",
    # Video
    "mp4", "webm", "avi", "mov", "mkv", "flv", "wmv", "gif",
    "m4v", "3gp", "ts", "mts", "vob", "ogv",
}

VIDEO_EXTENSIONS = {
    "mp4", "webm", "avi", "mov", "mkv", "flv", "wmv", "gif",
    "m4v", "3gp", "ts", "mts", "vob", "ogv",
}


@dataclass
class ParsedRequirement:
    id: str
    text: str
    source_file: str = ""
    line_number: int = 0
    # ``True`` when the requirement was synthesised from a bare URL by
    # ``_extract_url_requirement`` — it has no actual story content,
    # exists only so qa_persona's analyzer can detect the URL and crawl
    # the site. Downstream story generation must skip these entries.
    is_url_seed: bool = False


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_file(filepath: str, filename: str) -> tuple[list[str], str | None]:
    """Parse a file and return (list_of_text_lines, error_or_none)."""
    ext = filename.rsplit(".", 1)[1].lower() if "." in filename else ""

    if ext in ("txt", "md"):
        return _parse_text(filepath)
    elif ext == "docx":
        return _parse_docx(filepath)
    elif ext == "doc":
        return [], ".doc format is not supported directly. Please save the file as .docx"
    elif ext == "xlsx":
        return _parse_xlsx(filepath)
    elif ext == "csv":
        return _parse_csv(filepath)
    elif ext == "pdf":
        return _parse_pdf(filepath)
    elif ext in ("png", "jpg", "jpeg"):
        return _parse_image(filepath)
    elif ext in VIDEO_EXTENSIONS:
        return _parse_video(filepath, filename)
    else:
        return [], f"Unsupported format: .{ext}"


def _parse_text(filepath: str) -> tuple[list[str], str | None]:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            lines = [l.strip() for l in f if l.strip()]
        return lines, None
    except Exception as e:
        _logger.warning("read text file failed: %s", e)
        return [], f"Error reading file: {e}"


def _parse_docx(filepath: str) -> tuple[list[str], str | None]:
    if not HAS_DOCX:
        return [], "python-docx library is not installed. Run: pip install python-docx"
    try:
        doc = DocxDocument(filepath)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append(text)
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return lines, None
    except Exception as e:
        _logger.warning("read .docx failed: %s", e)
        return [], f"Error reading .docx: {e}"


def _parse_xlsx(filepath: str) -> tuple[list[str], str | None]:
    if not HAS_XLSX:
        return [], "openpyxl library is not installed. Run: pip install openpyxl"
    try:
        wb = load_workbook(filepath, read_only=True, data_only=True)
        lines = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    lines.append(" | ".join(cells))
        wb.close()
        return lines, None
    except Exception as e:
        _logger.warning("read .xlsx failed: %s", e)
        return [], f"Error reading .xlsx: {e}"


def _parse_csv(filepath: str) -> tuple[list[str], str | None]:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            lines = []
            for row in reader:
                cells = [c.strip() for c in row if c.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        return lines, None
    except Exception as e:
        _logger.warning("read .csv failed: %s", e)
        return [], f"Error reading .csv: {e}"


def _is_page_number(text: str) -> bool:
    """Detect lines that are just page numbers or page markers.

    Patterns detected:
      - Standalone numbers: "93", "116"
      - "Page 5", "page 12", "- 5 -", "— 5 —"
      - "5 / 20", "5/20" (page X of Y)
      - Roman numerals: "iv", "XII"
    """
    t = text.strip()

    # Pure number (1–9999)
    if re.fullmatch(r"\d{1,4}", t):
        return True

    # "Page N", "p. N", "Стр. N", "Сторінка N"
    if re.fullmatch(r"(?:page|p\.|стр\.?|сторінка)\s*\d{1,4}", t, re.IGNORECASE):
        return True

    # "- N -", "— N —"
    if re.fullmatch(r"[-—–]\s*\d{1,4}\s*[-—–]", t):
        return True

    # "N / M" or "N/M" (page of total)
    if re.fullmatch(r"\d{1,4}\s*/\s*\d{1,4}", t):
        return True

    # Roman numerals
    if re.fullmatch(r"[ivxlcdm]+", t, re.IGNORECASE) and len(t) <= 6:
        return True

    return False


def _strip_trailing_page_number(text: str) -> str:
    """Remove a trailing page number that got merged with content text.

    pypdf sometimes merges page numbers with adjacent content, producing
    lines like 'packet Submission 93' where 93 is the page number.

    Rules:
      - Only strip 1–4 digit numbers at the end of a line
      - Must be preceded by a space and a non-digit word
      - Don't strip if the number is clearly part of the text:
        * preceded by a version indicator (v, version, №, #)
        * preceded by a quantity word (up to, max, min, at least)
        * preceded by a unit or measurement context
    """
    # Pattern: text ending with space + 1-4 digits
    m = re.match(r"^(.+\b[a-zA-Zа-яА-ЯіІїЇєЄґҐ]{2,})\s+(\d{1,4})$", text.strip())
    if not m:
        return text

    body = m.group(1)

    # Don't strip if the number is contextually meaningful
    body_lower = body.lower()
    keep_patterns = [
        r"\bv(?:ersion)?\s*$",       # v2, version 2
        r"\b(?:№|#|no\.?)\s*$",      # №5, #5
        r"\b(?:up\s+to|max|min|at\s+least|менше|більше|до|від|мін|макс)\s*$",
        r"\b(?:step|крок|етап|phase|рівень|level)\s*$",
        r"\b(?:port|порт)\s*$",
        r"\b(?:error|code|код|помилка)\s*$",
        r"\b(?:item|елемент|пункт)\s*$",
        r"\b(?:\d+)\s*[-–—]\s*$",     # range like "10 - 20"
    ]
    for pat in keep_patterns:
        if re.search(pat, body_lower):
            return text

    return body


def _parse_pdf(filepath: str) -> tuple[list[str], str | None]:
    if not HAS_PDF:
        return [], "pypdf library is not installed. Run: pip install pypdf"
    try:
        reader = PdfReader(filepath)
        lines = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                for line in text.split("\n"):
                    stripped = line.strip()
                    if not stripped:
                        continue
                    # Skip standalone page numbers
                    if _is_page_number(stripped):
                        continue
                    # Strip trailing page numbers merged with content
                    stripped = _strip_trailing_page_number(stripped)
                    # Strip PDF underscore fill artifacts (e.g. "Field: _________")
                    stripped = re.sub(r"\s*_{2,}\s*", " ", stripped).strip()
                    stripped = re.sub(r"\s{2,}", " ", stripped)
                    # Skip if only underscores/whitespace remained
                    if not stripped or re.fullmatch(r"[_\s]+", stripped):
                        continue
                    lines.append(stripped)
        return lines, None
    except Exception as e:
        _logger.warning("read .pdf failed: %s", e)
        return [], f"Error reading .pdf: {e}"


def _parse_image(filepath: str) -> tuple[list[str], str | None]:
    if not HAS_PIL:
        return [], "Pillow library is not installed. Run: pip install Pillow"
    try:
        img = Image.open(filepath)
        info = f"[Image: {img.format}, {img.size[0]}x{img.size[1]}px, {img.mode}]"
        return [info], (
            "Image uploaded as attachment. For text recognition (OCR), "
            "we recommend using Tesseract OCR or Google Vision API separately. "
            "Image saved in the project for reference."
        )
    except Exception as e:
        _logger.warning("open image failed: %s", e)
        return [], f"Error opening image: {e}"


def _parse_video(filepath: str, filename: str) -> tuple[list[str], str | None]:
    """Parse video file — extract metadata (filename, size, format).

    Video files are treated as reference attachments. The filename and
    basic metadata are recorded so test cases can reference them.
    """
    try:
        file_size = os.path.getsize(filepath)
        ext = filename.rsplit(".", 1)[1].upper() if "." in filename else "VIDEO"

        if file_size < 1024:
            size_str = f"{file_size} B"
        elif file_size < 1024 * 1024:
            size_str = f"{file_size / 1024:.1f} KB"
        elif file_size < 1024 * 1024 * 1024:
            size_str = f"{file_size / (1024 * 1024):.1f} MB"
        else:
            size_str = f"{file_size / (1024 * 1024 * 1024):.2f} GB"

        info = f"[Video attachment: {filename} ({ext}, {size_str})]"
        return [info], None
    except Exception as e:
        _logger.warning("read video file failed: %s", e)
        return [], f"Error reading video file: {e}"


# ── Quality filtering helpers ────────────────────────────────────

# Conversational filler markers (English + Ukrainian/Russian)
_CONVERSATIONAL_PATTERNS = [
    # English filler
    r"\b(uh+|um+|hmm+|er+|ah+|oh+|yeah|nope|alright)\b",
    r"\b(you know|i mean|i guess|i think|let me|let's|by the way|kinda|sorta)\b",
    # Ukrainian filler / interjections
    r"\b(ну|ага|угу|тож|мабуть|дивись|слухай|розумієш|бачиш|знаєш|ось|от|отож|ото|оце|оцей|тут от|так от|от так|ну от|давай|давайте|короче|окей|гаразд|от|ото|там|отам)\b",
    # Russian filler
    r"\b(вот|это|типа|короче|значит|типо|как бы)\b",
]

# STRONG markers that indicate a formal requirement (modal verbs, capability
# verbs, explicit subject-verb patterns). These strongly suggest the line is
# a written requirement, not a conversational sentence that mentions a feature.
_STRONG_REQUIREMENT_MARKERS = [
    # Modal verbs (EN)
    r"\b(must|shall|should|needs?\s+to|has\s+to|have\s+to|ought\s+to|required\s+to)\b",
    # Capability verbs (EN)
    r"\b(allows?|enables?|supports?|provides?|permits?|lets?)\b",
    # Subject-verb patterns (EN)
    r"\b(the\s+)?(user|users|admin|administrator|customer|visitor|system|application|"
    r"app|platform|software|service|feature)\s+"
    r"(can|must|shall|should|will|may|is\s+able\s+to|should\s+be\s+able\s+to)\b",
    # Modal verbs (UA)
    r"\b(повинен|повинна|повинно|повинні|має|мають|потрібно|потрібен|необхідно|"
    r"мусить|мусять)\b",
    # Capability verbs (UA)
    r"\b(дозволяє|дозволяють|підтримує|надає|забезпечує)\b",
    # Subject-verb patterns (UA)
    r"\b(користувач|користувачі|адмін|адміністратор|відвідувач|клієнт|"
    r"система|застосунок|програма|сервіс|платформа)\s+"
    r"(може|можуть|повинен|повинна|повинні|має|мають|мусить|мусять)\b",
    # Imperative / instruction verbs common in requirement specs (EN)
    r"\b(allow\s+users?\s+to|support\s+(uploading|downloading|creating|editing|"
    r"deleting|viewing|filtering|searching|sorting)|implement|ensure\s+that)\b",
]

# WEAK markers — isolated feature keywords / UI nouns. Presence of only a weak
# marker is NOT enough to treat the line as a quality requirement, but it
# makes the line eligible for feature-keyword extraction.
_WEAK_FEATURE_MARKERS = [
    # CRUD/feature actions (EN)
    r"\b(login|log\s*in|sign\s*in|sign\s*up|register|logout|log\s*out|sign\s*out|"
    r"reset\s+password|search|filter|sort|upload|download|share|publish|subscribe|"
    r"checkout|purchase|profile|dashboard|settings|notifications)\b",
    # UI nouns (EN)
    r"\b(button|form|field|page|screen|dialog|modal|popup|dropdown|menu|tab|link|"
    r"section|panel|sidebar|navigation|header|footer)\b",
    # Features (UA)
    r"\b(вхід|вихід|реєстрація|пошук|фільтр|створення|редагування|видалення|"
    r"перегляд|оплата|кошик|оформлення|профіль|налаштування|сповіщення)\b",
    r"\b(кнопка|форма|сторінка|вікно|поле|меню|посилання|вкладка|екран|діалог)\b",
]

# Compiled once for speed
_CONV_RE = [re.compile(p, re.IGNORECASE) for p in _CONVERSATIONAL_PATTERNS]
_STRONG_REQ_RE = [re.compile(p, re.IGNORECASE) for p in _STRONG_REQUIREMENT_MARKERS]
_WEAK_FEAT_RE = [re.compile(p, re.IGNORECASE) for p in _WEAK_FEATURE_MARKERS]


def _is_metadata_line(text: str) -> bool:
    """Detect attachment/metadata markers like '[Video attachment: ...]'."""
    t = text.strip()
    if t.startswith("[") and t.endswith("]"):
        return True
    return False


# File extensions / structural tokens that should never become a feature.
# Without this filter the parser treats lines like "generation.py" or
# "requirements.txt" from a Markdown architecture diagram as standalone
# requirements, then generates a TC for "Test the generation.py page".
# That regression was reported as BUG-002…012, 014, 016, 017 in the
# self-audit (bug_reports_project.md).
_NON_FEATURE_EXTS = (
    ".py", ".pyc", ".js", ".ts", ".tsx", ".jsx", ".json", ".yml", ".yaml",
    ".toml", ".ini", ".cfg", ".conf", ".txt", ".log", ".lock", ".env",
    ".example", ".sample", ".md", ".rst", ".sh", ".bat", ".ps1",
    ".dockerfile", ".sql", ".html", ".css", ".scss", ".csv", ".tsv",
    ".xml", ".xlsx", ".docx", ".pdf",
)
_NON_FEATURE_TOKENS = {
    # bare filenames that occasionally appear without an extension
    "dockerfile", "makefile", "rakefile", "procfile", "license", "changelog",
    "readme", "manifest", "package", "gemfile", "pipfile", "go.sum", "go.mod",
}
_FILE_PATH_RE = re.compile(
    r"""^\s*[├└│─\s]*       # tree drawing characters from ``` diagrams
        ([A-Za-z0-9_./\\-]+ # path-like body
        \.[A-Za-z0-9]{1,8}) # extension
        (\s|$|/|\\)         # end-of-token boundary
    """, re.VERBOSE,
)
_TREE_DIAGRAM_RE = re.compile(r"[├└│─]")


def _is_non_feature_line(text: str) -> bool:
    """True for lines that look like file paths, code-tree diagrams,
    bare module names or other non-feature tokens.

    Examples this catches:
        ``generation.py``                  → file basename
        ``├── routes/``                    → tree-diagram leaf
        ``requirements.txt``               → tooling file
        ``.env.example``                   → dotfile sample
        ``app.py                ← Flask…`` → architecture-doc annotation
        ``- generation.py — Route module`` → bullet wrapping a file path
    """
    t = text.strip()
    if not t:
        return False

    # Tree drawing characters → architecture diagram, never a feature.
    if _TREE_DIAGRAM_RE.search(t):
        return True

    # Strip a single leading bullet / numbering marker so the path-detection
    # heuristic also catches "- generation.py — Route module".
    body = re.sub(r'^\s*(?:[-*\u2022]|\d+[\.)])\s+', '', t)

    # Whole-line / first-token match against file-extension list.
    head = body.split()[0] if body.split() else ""
    head_lc = head.lower().rstrip("/\\")
    if head_lc in _NON_FEATURE_TOKENS:
        return True
    for ext in _NON_FEATURE_EXTS:
        if head_lc.endswith(ext) and len(head) <= 60:
            return True

    # Inline path with arrow/comment ("app.py ← Flask application factory")
    # — first word is path-like with extension, rest is annotation.
    if _FILE_PATH_RE.match(body) and len(head) <= 60:
        return True

    # Numbered Markdown-style heading ("9. Версіонування та статистика" or
    # "## 9. Versioning and stats") — these are section labels, not
    # behaviours that can be tested. BUG-011 was caused by the heading
    # "Версіонування та статистика" being promoted to a TC.
    if re.match(r"^#{1,6}\s+\d+\.\s", t):
        return True

    return False


def _is_conversational(text: str) -> bool:
    """Detect conversational/transcript fragments that should NOT be requirements."""
    lower = text.lower().strip()

    # Empty or very short
    if len(lower) < 15:
        return True

    # Question fragments (often conversational, especially short ones)
    if lower.endswith("?") and len(lower.split()) < 12:
        return True

    # Truncated mid-sentence
    if lower.endswith("...") or lower.endswith("…"):
        return True

    # Count TOTAL filler matches across all patterns (not distinct patterns)
    filler_hits = 0
    for pat in _CONV_RE:
        filler_hits += len(pat.findall(lower))
        if filler_hits >= 2:
            return True

    # First-person narrative ("я + verb") without requirement signal = conversational
    if (re.search(r"\b(я|мені|мене|мій|моя|моє|мої)\b", lower) and
            not _has_requirement_signal(lower)):
        return True

    # Direct second-person address (ти/ви) + short sentence + no requirement signal
    if (re.search(r"\b(ти|ви|тебе|тобі|вас|вам|твій|твоя|твоє|твої)\b", lower) and
            len(lower.split()) < 18 and
            not _has_requirement_signal(lower)):
        return True

    # Personal name (capitalized non-initial word that's a common Slavic name)
    if re.search(
        r"\b(Юра|Юрій|Юри|Саша|Саші|Оля|Олі|Марія|Марії|Іван|Петро|"
        r"Тарас|Олег|Сергій|Сергія|Анна|Анни|Роман|Романа|Володимир|"
        r"Микола|Миколи|Степан|Андрій|Андрія|Віктор|Вікторія)\b",
        text,
    ):
        return True

    return False


def _has_strong_requirement_signal(text: str) -> bool:
    """Check if text has an explicit modal verb, capability verb, or subject-verb
    pattern that marks it as a formally written requirement."""
    for pat in _STRONG_REQ_RE:
        if pat.search(text):
            return True
    return False


def _has_weak_feature_signal(text: str) -> bool:
    """Check if text contains any recognizable feature keyword or UI noun."""
    for pat in _WEAK_FEAT_RE:
        if pat.search(text):
            return True
    return False


# Backwards-compat alias
def _has_requirement_signal(text: str) -> bool:
    return _has_strong_requirement_signal(text) or _has_weak_feature_signal(text)


def _is_quality_requirement(text: str) -> bool:
    """Strict check: does this text look like a formally written requirement?

    Requires a STRONG signal (modal verb / subject-verb / capability verb).
    Lines that only mention feature keywords casually are not quality — they
    must be normalized via feature extraction instead.
    """
    t = text.strip()

    if len(t) < 15 or len(t) > 600:
        return False

    if _is_metadata_line(t):
        return False

    if _is_conversational(t):
        return False

    return _has_strong_requirement_signal(t)


# ── Feature extraction from corpus (transcript fallback) ─────────

# UI element patterns: "login button", "search form", etc.
# Only match ONE modifier word to avoid dragging conversational context.
_UI_FEATURE_RE = re.compile(
    r"\b([a-zа-яёіїє]{3,})\s+"
    r"(button|form|page|screen|field|input|menu|dialog|modal|popup|dropdown|tab|link|"
    r"кнопка|форма|сторінка|вікно|поле|меню|посилання|вкладка|екран|діалог)\b",
    re.IGNORECASE,
)

# Words that should NOT be treated as feature modifiers (pronouns, articles, fillers)
_UI_MODIFIER_BLACKLIST = {
    # English
    "the", "this", "that", "these", "those", "some", "any", "every", "each",
    "one", "two", "three", "new", "old", "first", "last", "other", "another",
    "following", "current", "main", "next", "previous", "same", "my", "your",
    "our", "their", "his", "her", "its",
    # Ukrainian pronouns / articles / common noise
    "цей", "ця", "це", "ці", "той", "та", "ті", "тут", "там", "так",
    "наш", "наша", "наше", "мій", "моя", "моє", "ваш", "ваша", "ваше",
    "перший", "останній", "новий", "старий", "інший", "інша",
    "наступний", "попередній", "головний", "головна", "поточний",
    # Russian
    "этот", "эта", "это", "эти", "тот", "та", "те", "наш",
}

# URL detection
_URL_RE = re.compile(
    r"(https?://)?([a-z0-9][a-z0-9\-]*\.)+[a-z]{2,}(/[^\s]*)?",
    re.IGNORECASE,
)

# Discrete feature keywords. Each keyword maps to a (feature_id, description).
# feature_id is used for global deduplication so English and Ukrainian
# synonyms for the same feature (e.g. "login" / "увійти") produce a single
# normalized requirement.
_FEATURE_KEYWORDS: dict[str, tuple[str, str]] = {
    # Auth
    "login":              ("auth.login",        "User can log in"),
    "log in":             ("auth.login",        "User can log in"),
    "sign in":            ("auth.login",        "User can sign in"),
    "вхід":               ("auth.login",        "User can log in"),
    "увійти":             ("auth.login",        "User can log in"),
    "logout":             ("auth.logout",       "User can log out"),
    "log out":            ("auth.logout",       "User can log out"),
    "sign out":           ("auth.logout",       "User can sign out"),
    "вихід":              ("auth.logout",       "User can log out"),
    "вийти":              ("auth.logout",       "User can log out"),
    "register":           ("auth.register",     "User can register an account"),
    "registration":       ("auth.register",     "User can register an account"),
    "sign up":            ("auth.register",     "User can sign up"),
    "реєстрація":         ("auth.register",     "User can register an account"),
    "реєстрації":         ("auth.register",     "User can register an account"),
    "зареєструватися":    ("auth.register",     "User can register an account"),
    "password reset":     ("auth.pwd_reset",    "User can reset password"),
    "forgot password":    ("auth.pwd_reset",    "User can recover forgotten password"),
    "reset password":     ("auth.pwd_reset",    "User can reset password"),
    "відновлення паролю": ("auth.pwd_reset",    "User can reset password"),
    "two-factor":         ("auth.2fa",          "Two-factor authentication is supported"),
    "2fa":                ("auth.2fa",          "Two-factor authentication is supported"),
    # Search / Filter / Sort
    "search":             ("search",            "User can search content"),
    "пошук":              ("search",            "User can search content"),
    "filter":             ("filter",            "User can filter results"),
    "фільтр":             ("filter",            "User can filter results"),
    "sort":               ("sort",              "User can sort results"),
    "сортування":         ("sort",              "User can sort results"),
    # CRUD
    "create account":     ("crud.create_acct",  "User can create an account"),
    "edit profile":       ("crud.edit_profile", "User can edit profile"),
    "delete account":     ("crud.del_acct",     "User can delete account"),
    "створити":           ("crud.create",       "User can create items"),
    "редагувати":         ("crud.edit",         "User can edit items"),
    "видалити":           ("crud.delete",       "User can delete items"),
    # Commerce
    "checkout":           ("commerce.checkout", "User can complete checkout"),
    "оформлення":         ("commerce.checkout", "User can complete checkout"),
    "shopping cart":      ("commerce.cart",     "Shopping cart functionality works"),
    "add to cart":        ("commerce.cart",     "User can add items to cart"),
    "кошик":              ("commerce.cart",     "Shopping cart functionality works"),
    "payment":             ("commerce.payment",  "User can make a payment"),
    "оплата":              ("commerce.payment",  "User can make a payment"),
    "purchase":            ("commerce.purchase", "User can purchase items"),
    # Content
    "upload":             ("content.upload",    "User can upload files"),
    "завантажити":        ("content.upload",    "User can upload files"),
    "download":           ("content.download",  "User can download files"),
    "share":              ("content.share",     "User can share content"),
    "comment":            ("content.comment",   "User can leave comments"),
    "subscribe":          ("content.subscribe", "User can subscribe"),
    # Navigation / Account
    "dashboard":          ("nav.dashboard",     "Dashboard renders every configured widget with data"),
    "settings":           ("nav.settings",      "User can change settings"),
    "налаштування":       ("nav.settings",      "User can change settings"),
    "profile":            ("nav.profile",       "User can view profile"),
    "профіль":            ("nav.profile",       "User can view profile"),
    "notifications":      ("nav.notifications", "User receives notifications"),
    "сповіщення":         ("nav.notifications", "User receives notifications"),
}


def _extract_url_requirement(line: str) -> str | None:
    """If a line contains a URL, return a clean 'domain: path' requirement
    describing what should be tested. Handles lines like:

      'Create a checklist for the following page: https://testfort.com/abc'
      → 'testfort.com: abc'
    """
    m = _URL_RE.search(line)
    if not m:
        return None
    matched = m.group(0)
    # Re-run to grab the groups properly (we want domain and path separately)
    inner = re.match(
        r"(https?://)?([^/\s]+\.[a-z]{2,})(/[^\s]*)?", matched, re.IGNORECASE,
    )
    if not inner:
        return None
    domain = inner.group(2)
    path = (inner.group(3) or "").strip("/")
    path = path.replace("/", " > ").replace("-", " ").replace("_", " ")
    if path:
        return f"{domain}: {path}"
    return domain


def _extract_feature_requirements(line: str) -> list[tuple[str, str]]:
    """Extract feature-keyword-based requirements from a single line.

    Used when a line is conversational or lacks strong requirement signals
    but contains recognizable feature mentions (e.g. 'there is a login
    button' → 'User can log in'). Returns list of (feature_id, description)
    tuples so the caller can deduplicate across lines.
    """
    lower = line.lower()
    found: list[tuple[str, str]] = []
    seen_ids: set[str] = set()

    # Longer keywords match first so 'sign up' beats 'sign'
    for kw in sorted(_FEATURE_KEYWORDS.keys(), key=len, reverse=True):
        if re.search(rf"(?<![\w\-]){re.escape(kw)}(?![\w\-])", lower):
            feat_id, req = _FEATURE_KEYWORDS[kw]
            if feat_id not in seen_ids:
                seen_ids.add(feat_id)
                found.append((feat_id, req))
    return found


def _extract_features_from_corpus(lines: list[str]) -> list[str]:
    """Extract feature mentions from long-form text (e.g., a video transcript).

    When all the user provides is unstructured prose, scan it for feature
    mentions and return one normalized requirement per detected feature.
    """
    full = " ".join(lines).lower()
    if not full:
        return []

    found: dict[str, str] = {}

    # 1. UI element patterns: "search button", "login form"
    for m in _UI_FEATURE_RE.finditer(full):
        modifier = m.group(1).strip().lower()
        ui = m.group(2).strip().lower()
        if modifier in _UI_MODIFIER_BLACKLIST:
            continue
        key = f"{modifier} {ui}"
        if key not in found:
            found[key] = f"User can use the {modifier} {ui}"

    # 2. Discrete feature keywords, deduplicated by feature_id
    seen_ids: set[str] = set()
    for kw in sorted(_FEATURE_KEYWORDS.keys(), key=len, reverse=True):
        if re.search(rf"(?<![\w\-]){re.escape(kw)}(?![\w\-])", full):
            feat_id, requirement = _FEATURE_KEYWORDS[kw]
            if feat_id in seen_ids:
                continue
            seen_ids.add(feat_id)
            if requirement not in found:
                found[requirement] = requirement

    return list(found.values())


# ── Requirement Splitter ─────────────────────────────────────────

def split_into_requirements(lines: list[str]) -> list[ParsedRequirement]:
    """Split raw lines into structured ParsedRequirement objects.

    Multi-pass strategy:
      1. Extract structured items (REQ-IDs, prefixed checklist IDs, numbered, bullets)
      2. Extract quality requirement-like sentences (with feature signals)
      3. If nothing solid found, extract feature mentions from the corpus
      4. Last-resort fallback so we always return something usable

    Conversational text and metadata markers are filtered out.
    """
    counter = 1
    structured: list[ParsedRequirement] = []
    plain_candidates: list[tuple[int, str]] = []
    seen_urls: set[str] = set()

    # Compiled patterns for structured extraction
    req_id_re = re.compile(r"^(REQ-\d+)\s*[:\.\-]?\s*(.+)", re.IGNORECASE)
    # TestFort-style checklist IDs: AUTH_001, SRCH_001, etc.
    cl_id_re = re.compile(r"^([A-Z]{2,6}_\d+)\s*[:\.\-]?\s*(.+)")
    # Numbered: 1. ... or 1) ...
    numbered_re = re.compile(r"^\d+[\.\)]\s+(.+)")
    # Bullet: - ... or * ... or • ...
    bullet_re = re.compile(r"^[\-\*\u2022]\s+(.+)")

    for i, line in enumerate(lines):
        text = line.strip()
        if not text:
            continue

        # Skip metadata markers
        if _is_metadata_line(text):
            continue

        # Skip file paths, tree diagrams, code-block leftovers, numbered
        # Markdown headings — these are not testable behaviours. Run
        # BEFORE URL extraction so "generation.py" doesn't get matched
        # as a domain by the loose URL regex.
        if _is_non_feature_line(text):
            continue

        # URL detection — emit a clean domain+path requirement for any line
        # that contains a URL (e.g. an instruction to generate tests for a
        # given web page).
        url_req = _extract_url_requirement(text)
        if url_req and url_req not in seen_urls:
            seen_urls.add(url_req)
            structured.append(ParsedRequirement(
                id=f"REQ-{counter:03d}",
                text=f"Test the {url_req} page",
                line_number=i + 1,
                is_url_seed=True,
            ))
            counter += 1
            continue

        # Try REQ-ID format
        m = req_id_re.match(text)
        if m:
            req_text = m.group(2).strip()
            if len(req_text) >= 5:
                structured.append(ParsedRequirement(
                    id=m.group(1).upper(), text=req_text, line_number=i + 1,
                ))
                continue

        # Try TestFort checklist ID format (AUTH_001, etc.)
        m = cl_id_re.match(text)
        if m:
            req_text = m.group(2).strip()
            if len(req_text) >= 5:
                structured.append(ParsedRequirement(
                    id=m.group(1), text=req_text, line_number=i + 1,
                ))
                continue

        # Try numbered list
        m = numbered_re.match(text)
        if m:
            req_text = m.group(1).strip()
            if (len(req_text) >= 5 and not _is_conversational(req_text)
                    and not _is_non_feature_line(req_text)):
                structured.append(ParsedRequirement(
                    id=f"REQ-{counter:03d}", text=req_text, line_number=i + 1,
                ))
                counter += 1
                continue

        # Try bullet list
        m = bullet_re.match(text)
        if m:
            req_text = m.group(1).strip()
            if (len(req_text) >= 5 and not _is_conversational(req_text)
                    and not _is_non_feature_line(req_text)):
                structured.append(ParsedRequirement(
                    id=f"REQ-{counter:03d}", text=req_text, line_number=i + 1,
                ))
                counter += 1
                continue

        # Plain line — collect for quality analysis later
        plain_candidates.append((i + 1, text))

    # Phase 2: Quality-filter plain candidates OR extract features from
    # conversational lines that mention known feature keywords.
    quality: list[ParsedRequirement] = []
    seen_feat_ids: set[str] = set()

    for line_num, text in plain_candidates:
        if _is_quality_requirement(text):
            quality.append(ParsedRequirement(
                id=f"REQ-{counter:03d}", text=text, line_number=line_num,
            ))
            counter += 1
            continue

        # Not a quality line — but if it mentions recognizable features,
        # emit normalized requirements for each (globally deduplicated
        # by feature_id so EN/UA synonyms collapse into one).
        for feat_id, feat_text in _extract_feature_requirements(text):
            if feat_id in seen_feat_ids:
                continue
            seen_feat_ids.add(feat_id)
            quality.append(ParsedRequirement(
                id=f"REQ-{counter:03d}", text=feat_text, line_number=line_num,
            ))
            counter += 1

    # If we have decent results from structured + quality, return them
    if structured or quality:
        return structured + quality

    # Phase 3: Corpus extraction (transcript fallback — no per-line matches).
    corpus_lines = [t for _, t in plain_candidates]
    features = _extract_features_from_corpus(corpus_lines)
    if features:
        return [
            ParsedRequirement(id=f"REQ-{idx:03d}", text=feat, line_number=0)
            for idx, feat in enumerate(features, start=1)
        ]

    # Phase 4: No testable input found. Return an empty list so the caller
    # can surface a clear error to the user instead of generating garbage.
    return []
